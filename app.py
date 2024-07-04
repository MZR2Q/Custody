from flask import Flask, render_template, request, redirect, send_file, session,jsonify
import sqlite3
import datetime
import html
from docx import Document
import sqlite3
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
import hashlib
import pandas as pd
from mailer import Mailer
from io import BytesIO
from barcode import get_barcode, generate 
from barcode.writer import ImageWriter
from sender import sender
today = datetime.date.today().day
import base64
 

 
app = Flask(__name__)
app.secret_key = '66745674564564564569'

def export_to_excel(table_name):
    connection = sqlite3.connect('database.db')
    query = f"SELECT * FROM {table_name}"
    df = pd.read_sql_query(query, connection)
    
    # مسار حفظ الملف بجوار ملف app.py
    file_path = os.path.join(os.path.dirname(__file__), f"{table_name}_data.xlsx")
    
    df.to_excel(file_path, index=False)
    connection.close()
    return file_path

def encrippt(prs):

    p1e = hashlib.md5(prs.encode()).hexdigest()
    p1d1 = p1e.translate({ord('b'): None})
    p1d2 = p1d1.translate({ord('8'): None})
    p1 = p1d2[0:9]

    p2e = hashlib.sha1(prs.encode()).hexdigest()
    p2d1 = p2e.translate({ord('e'): None})
    p2d2 = p2d1.translate({ord('6'): None})
    p2 = p2d2[0:6]

    p3e = hashlib.sha256(prs.encode()).hexdigest()
    p3d1 = p3e.translate({ord('1'): None})
    p3d2 = p3d1.translate({ord('l'): None})
    p3 = p3d2[0:10]

    return str(p1)+str(p2)+str(p3)



@app.route('/export_data_data')
def export_data_data():
    try:
        file_path = export_to_excel('data')
        with open(file_path, 'rb') as file:
            file_data = base64.b64encode(file.read()).decode('utf-8')
        return jsonify({'file_data': file_data})
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/export_data_storage')
def export_data_storage():
    try:
        file_path = export_to_excel('storage')
        with open(file_path, 'rb') as file:
            file_data = base64.b64encode(file.read()).decode('utf-8')
        return {'file_data': file_data}
    except Exception as e:
        return {'error': str(e)}

@app.route('/Dashboard/export_data_Device_of_coll')
def export_data_Device_of_coll():
    try:
        file_path = export_to_excel('Device_of_coll')
        with open(file_path, 'rb') as file:
            file_data = base64.b64encode(file.read()).decode('utf-8')
        return jsonify({'file_data': file_data})
    except Exception as e:
        return jsonify({'error': str(e)})












@app.route('/', methods =["GET", "POST"] )
def loginn():
    db = sqlite3.connect("AcD.db")
    mydb = db.cursor()
    if request.method == 'POST':
        user = html.escape(request.form.get('emailUsername'))
        passe = html.escape(request.form.get('password'))

        query = "SELECT * FROM usersz WHERE Email = ? AND Password = ?"
        mydb.execute(query, (user, encrippt(passe)))
        AcountValues = mydb.fetchone()
        if AcountValues == None:
            return render_template('login.html', errormassge='Errore User Or password')
        else:
            session['UserEmail'] = AcountValues[2]
            session['User'] = AcountValues[1]
            session['type'] = AcountValues[4]
            return redirect('/Dashboard')
    return render_template('login.html')


@app.route('/Dashboard/acountmang', methods=['GET', 'POST'])
def manage_accounts():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['Tic', 'all']:
        return redirect('/')
    if  session.get('UserEmail') !=  'Admin@example.com':
        return redirect('/')




    if request.method == 'POST':
        user_id = request.form.get('user_id')
        new_name = request.form.get('new_name')
        new_email = request.form.get('new_email')
        new_password = request.form.get('new_password')

        # Hash the new password before storing it in the database

        conn = sqlite3.connect('AcD.db')
        cursor = conn.cursor()
        cursor.execute('''
            UPDATE usersz
            SET Name=?, Email=?, Password=?
            WHERE i=?
        ''', (new_name, new_email, encrippt(new_password), user_id))
        conn.commit()
        conn.close()
        return redirect('/Dashboard/acountmang')
    conn = sqlite3.connect('AcD.db')
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM usersz')
    users = cursor.fetchall()
    conn.close()

    return render_template('AcountMangment.html', userEmail=session.get('UserEmail') , userName=session.get('User'), users=users, type=session.get('type'))


@app.route('/signup', methods =["GET", "POST"])
def signup():
    if  session.get('UserEmail') !=  'Admin@example.com':
        return redirect('/')
    db = sqlite3.connect("AcD.db")
    mydb = db.cursor()
    if request.method == "POST":
        _Name = html.escape(request.form.get('fullName'))
        _Email = html.escape(request.form.get('email'))
        _password = html.escape(request.form.get('password'))
        _type = html.escape(request.form.get('type'))
        query = "INSERT INTO usersz (Name, Email, Password, type) VALUES (?, ?, ?, ?)"
        mydb.execute(query, (_Name, _Email, encrippt(_password), _type))
        db.commit()
        db.close()

        massge = f'https://labmanagment.com:8080/ <br> You can now access UHB custody <br>  User : {_Email} <br>  Password : {_password}'
        sender(_Email,massge)
        return redirect('/Dashboard/devicemange')























@app.route('/send/<labname>/<pcname>', methods=['GET', 'POST'])
def home(labname,pcname):
    if request.method == 'POST':
        room_id = labname
        pc_id = pcname
        date = today
        description = html.escape( request.form['description'])
        student_id = html.escape(request.form['student_id'])
        studname = html.escape(request.form['StudName'])


        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("INSERT INTO data (RoomId, Discription, PcId, Date, StudId, StudName) VALUES (?, ?, ?, ?, ?, ?)",
                        (room_id, description, pc_id, date, student_id, studname))
        conn.commit()
        conn.close()

        return render_template('sender.html')

    return render_template('sender.html', labn=labname, pcn = pcname, type=session.get('type'))



@app.route('/admm', methods=['GET', 'POST'])
def admm():
    if request.method == 'POST':
        search_res =  html.escape(request.form.get('sea'))
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM storage WHERE LabName=?",(str(search_res),))
        result = cursor.fetchall()
        print(result)
        return render_template('addDev.html', userEmail=session.get('UserEmail') , userName=session.get('User'), result=result , type=session.get('type'))

    return render_template('addDev.html', userEmail=session.get('UserEmail') , userName=session.get('User'), type=session.get('type'))





@app.route('/add_devices', methods=['POST'])
def add_device():
    if request.method == 'POST':
        lab_name =  html.escape(request.form['lab_name'])
        pc_num =  html.escape(request.form['pc_num'])
        pc_prosser =  html.escape(request.form['pc_prosser'])
        pc_model =  html.escape(request.form['pc_model'])
        pc_keyboard =  html.escape(request.form['pc_keyboard'])
        pc_mouse =  html.escape(request.form['pc_mouse'])
        details =  html.escape(request.form['details'])
        
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        c.execute("INSERT INTO storage (LabName, PcNum, PcProsser, PcModel, PcKybord, PcMouse, Details) VALUES (?, ?, ?, ?, ?, ?, ?)",
                  (lab_name, pc_num, pc_prosser, pc_model, pc_keyboard, pc_mouse, details))
        conn.commit()
        conn.close()

        return redirect('/Dashboard/devicemange?Categ=Add')






@app.route('/crea', methods=['POST'])
def crea():
    if request.method == 'POST':
        labnum =  html.escape(request.form.get('labNum'))
        doc = Document()

        title = doc.add_paragraph(f'Report For Lab : {labnum}')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER  

        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        headers = [ 'Lab Room', 'Number', 'Pressor', 'Model', 'Keyboard', 'Mouse', 'Details']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header

        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("SELECT  LabName, PcNum, PcProsser, PcModel, PcKybord, PcMouse, Details FROM storage WHERE labName=(?)",(str(labnum),))
        data = cursor.fetchall()

        for row in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        doc.save(f'Lab_{labnum}.docx')
        p = f'Lab_{labnum}.docx'
        conn.close()
        return send_file(p)

    return redirect('/admm')









@app.route('/Dashboard')
def Admin():
    if session.get('UserEmail') == None:
        return redirect('/')
    return render_template('index.html', userEmail=session.get('UserEmail') , userName=session.get('User') , type=session.get('type'))



@app.route('/Dashboard/devicemange')
def Mange():
    print(session.get('type'))
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['Tic', 'all']:
        return redirect('/')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("SELECT count(*) FROM storage")
    all_lab_dev = cursor.fetchall()
    print(all_lab_dev[0][0])
    cursor.execute("SELECT count(*) FROM Device_of_coll_New")
    all_emp_dev = cursor.fetchall()

    cursor.execute("SELECT count(*) FROM data")
    all_req_dev = cursor.fetchall()

    categ = request.args.get('Categ')  
    if categ == 'labdev':
        print('labdev')
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM storage")
        allLabDev = cursor.fetchall()
        return render_template('searchval.html', userEmail=session.get('UserEmail') , userName=session.get('User'), allLabDev=allLabDev,categ=categ, all_lab_dev=all_lab_dev[0][0],all_emp_dev=all_emp_dev[0][0],all_req_dev=all_req_dev[0][0] , type=session.get('type'))
    elif categ == 'requ':
        print('requ')
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM data")
        allLabDev = cursor.fetchall()
        return render_template('searchval.html', userEmail=session.get('UserEmail') , userName=session.get('User'), allLabDev=allLabDev,categ=categ, all_lab_dev=all_lab_dev[0][0],all_emp_dev=all_emp_dev[0][0],all_req_dev=all_req_dev[0][0] , type=session.get('type'))
    elif categ == 'coldev':
        print('coldev')
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM Device_of_coll_New")
        allLabDev = cursor.fetchall()
        return render_template('searchval.html', userEmail=session.get('UserEmail') , userName=session.get('User'), allLabDev=allLabDev,categ=categ, all_lab_dev=all_lab_dev[0][0],all_emp_dev=all_emp_dev[0][0],all_req_dev=all_req_dev[0][0] , type=session.get('type'))
    elif categ == 'Add':
        print('coldev')
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM Device_of_coll_New")
        allLabDev = cursor.fetchall()
        return render_template('AddDevice.html', userEmail=session.get('UserEmail') , userName=session.get('User'), allLabDev=allLabDev,categ=categ, all_lab_dev=all_lab_dev[0][0],all_emp_dev=all_emp_dev[0][0],all_req_dev=all_req_dev[0][0] , type=session.get('type'))
    
    return render_template('devicesMange.html', userEmail=session.get('UserEmail') , userName=session.get('User'), all_lab_dev=all_lab_dev[0][0],all_emp_dev=all_emp_dev[0][0],all_req_dev=all_req_dev[0][0], type=session.get('type'))


@app.route('/deletd', methods=['GET', 'POST'])
def deleted():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['Tic', 'all']:
        return redirect('/')
    if request.method == 'POST':
        request_id = request.form['request_id']
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("DELETE FROM storage WHERE id=?", (request_id,))
        conn.commit()
        conn.close()
        print(request_id)

    return redirect('/Dashboard/devicemange?Categ=labdev')














@app.route('/generate_qr/<path:url>')
def generate_qr(url):
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['Tic', 'all']:
        return redirect('/')

    buffer = BytesIO()
    generate('code128', url, writer=ImageWriter(), output=buffer)
    buffer.seek(0)

    return send_file(buffer, mimetype='image/png', as_attachment=True, download_name='qrcode.png')

@app.route('/Dashboard/Scanner')
def scan_barcode():
    return render_template('Scanner.html', userEmail=session.get('UserEmail') , userName=session.get('User'), type=session.get('type'))





















@app.route('/process_barcode', methods=['POST'])
def process_barcode():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['Tic', 'all']:
        return redirect('/')
    barcode_value = request.json['barcode']

    product_name = get_product_name(barcode_value) 

    return jsonify(product_name)



def get_product_name(barcode_value):
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    



    if barcode_value[:2] == 'ED':
        print(barcode_value,'xxxxx')
        cursor.execute('SELECT * FROM Device_of_coll_New WHERE id=?',(barcode_value[3:],))
        date = cursor.fetchone()
        info = {    
            'type' : 'Employ Device',
            'EmpName':f"Employ Name : {date[1]}",
            'EmpEmpName':f"Job title : {date[2]}",
            'EmpNum':f"Employ Number : {date[3]}",
            'Mange':f"Department  : {date[4]}",
            'DeviceType':f"Device Type : {date[5]}",
            'manufacturingCompanies':f"MC : {date[6]}",
            'modale':f"modal : {date[7]}",
            'serialnum':f"Serial Number : {date[9]}",
            'location':f"Location : {date[9]}",
            'offecenum':f"Office Num : {date[10]}"

        }
        print(info)
        return info

    elif barcode_value[:2] == 'LD':
        print(barcode_value,'xxxxx')
        cursor.execute('SELECT * FROM storage WHERE id=?',(barcode_value[3:],))
        date = cursor.fetchone()

        info = {    
            'type' : 'Lab Device',
            'PCID':date[0],
            'Labname':date[2],
            'PcNum':date[3],
            'PcProsore':date[4],
            'Pcmodel':date[5],
            'pckybord':date[6],
            'Pcmouse':date[7],
            'details':date[8],

        }
        return info

    



























@app.route('/insert', methods=['POST'])
def insert():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['Tic', 'all']:
        return redirect('/')
    empName = request.form['empName']
    empempName = request.form['empempName']
    empNum = request.form['empNum']
    mange = request.form['mange']
    deviceType = request.form['deviceType']
    manufacturingCompanies = request.form['manufacturingCompanies']
    modale = request.form['modale']
    serialNum = request.form['serialNum']
    location = request.form['location']
    officeNum = request.form['officeNum']

    # Connect to the SQLite database
    connection = sqlite3.connect('database.db')
    cursor = connection.cursor()

    # Insert data into the table
    cursor.execute("INSERT INTO Device_of_coll_New (empName, empempName, empNum, mange, deviceType, manufacturingCompanies, modale, serialNum, location, officeNum) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                   (empName, empempName, empNum, mange, deviceType, manufacturingCompanies, modale, serialNum, location, officeNum))

    # Commit the transaction and close the connection
    connection.commit()
    connection.close()

    return redirect('/Dashboard/devicemange?Categ=Add')




@app.route('/logout')
def Leav():
    session.clear()
    return redirect('/')





@app.route('/add_item', methods=['POST'])
def add_item():
        # List of all field names
    if request.method == 'POST':
        officeNumber = html.escape(str(request.form.get('officeNumber')))
        empname = html.escape(str(request.form.get('name')))
        empnum = html.escape(str(request.form.get('empid')))
        desk = html.escape(request.form.get("desk"))
        desk_accessory = html.escape(request.form.get("desk_accessory"))
        office_cabinet = html.escape(request.form.get("office_cabinet"))
        mobile_drawers = html.escape(request.form.get("mobile_drawers"))
        office_tools = html.escape(request.form.get("office_tools"))
        small_tea_table = html.escape(request.form.get("small_tea_table"))
        large_tea_table = html.escape(request.form.get("large_tea_table"))
        swivel_chair = html.escape(request.form.get("swivel_chair"))
        visitor_chair = html.escape(request.form.get("visitor_chair"))
        leather_sofa = html.escape(request.form.get("leather_sofa"))
        double_sofa = html.escape(request.form.get("double_sofa"))
        triple_sofa = html.escape(request.form.get("triple_sofa"))
        electric_stapler = html.escape(request.form.get("electric_stapler"))
        paper_shredder = html.escape(request.form.get("paper_shredder"))
        hole_punch_large = html.escape(request.form.get("hole_punch_large"))
        metal_cabinet = html.escape(request.form.get("metal_cabinet"))
        safe_small = html.escape(request.form.get("safe_small"))
        safe_large = html.escape(request.form.get("safe_large"))
        computer = html.escape(request.form.get("computer"))
        extra_monitor = html.escape(request.form.get("extra_monitor"))
        printer_plain = html.escape(request.form.get("printer_plain"))
        printer_color = html.escape(request.form.get("printer_color"))
        scanner = html.escape(request.form.get("scanner"))
        copier = html.escape(request.form.get("copier"))
        calculator = html.escape(request.form.get("calculator"))
        phone_regular = html.escape(request.form.get("phone_regular"))
        phone_digital = html.escape(request.form.get("phone_digital"))
        fields = [
        'office_number' ,
        'name' ,
        'empid' ,
        'desk',
        'desk_accessory',
        'office_cabinet' ,
        'mobile_drawers' ,
        'office_tools' ,
        'small_tea_table' ,
        'large_tea_table' ,
        'swivel_chair' ,
        'visitor_chair' ,
        'leather_sofa' ,
        'double_sofa' ,
        'triple_sofa' ,
        'electric_stapler' ,
        'paper_shredder' ,
        'hole_punch_large' ,
        'metal_cabinet' ,
        'safe_small' ,
        'safe_large' ,
        'computer' ,
        'extra_monitor' ,
        'printer_plain' ,
        'printer_color' ,
        'scanner' ,
        'copier' ,
        'calculator' ,
        'phone_regular' ,
        'phone_digital' 
        ]



        connection = sqlite3.connect('office_inventory.db')
        cursor = connection.cursor()

        # Prepare the SQL query dynamically based on the fields
        query = f'''
            INSERT INTO office ({', '.join(fields)})
            VALUES ("{officeNumber}","{empname}","{empnum}","{desk}","{desk_accessory}","{office_cabinet}","{mobile_drawers}","{office_tools}","{small_tea_table}","{large_tea_table}","{swivel_chair}","{visitor_chair}","{leather_sofa}","{double_sofa}","{triple_sofa}","{electric_stapler}","{paper_shredder}","{hole_punch_large}","{metal_cabinet}","{safe_small}","{safe_large}","{computer}","{extra_monitor}","{printer_plain}","{printer_color}","{scanner}","{copier}","{calculator}","{phone_regular}","{phone_digital}")
        '''
        print({', '.join(fields)})
        print(request.form.get('desk'))
        cursor.execute(query)

        connection.commit()
        connection.close()

        return redirect('/Dashboard/furniture/add')



@app.route('/Dashboard/furniture/add')
def addfurniture():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['furniture', 'all']:
        return redirect('/')
    fields = [
        {'name': 'desk', 'label': 'طاولة مكتب'},
    ]

    return render_template('furnitureadd.html', fields = fields, userEmail=session.get('UserEmail') , userName=session.get('User') , type=session.get('type'))



@app.route('/Dashboard/furniture')
def furniture():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['furniture', 'all']:
        return redirect('/')
    connection = sqlite3.connect('office_inventory.db')
    cursor = connection.cursor()
    cursor.execute('SELECT name, desk FROM office')
    all = cursor.fetchall()
    print(all)
    x = 0
    for i in all:
        x += int(i[1])
    connection = sqlite3.connect('office_inventory.db')
    cursor = connection.cursor()
    cursor.execute('SELECT * FROM office')
    ttcc = cursor.fetchall()
    dd = 0
    for i in ttcc:
        dd += 1
    return render_template('furnituremang.html',x=x,dd=dd, userEmail=session.get('UserEmail') , userName=session.get('User') , type=session.get('type'))



@app.route('/Dashboard/furniture/disks')
def diskfur():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['furniture', 'all']:
        return redirect('/')
    connection = sqlite3.connect('office_inventory.db')
    cursor = connection.cursor()
    cursor.execute('SELECT name, desk FROM office')
    all = cursor.fetchall()
    print(all)
    x = 0
    for i in all:
        x += int(i[1])
    connection = sqlite3.connect('office_inventory.db')
    cursor = connection.cursor()
    cursor.execute('SELECT * FROM office')
    ttcc = cursor.fetchall()
    dd = 0
    for i in ttcc:
        dd += 1
    return render_template('fur_Disk.html',all=all,x=x,dd=dd , userEmail=session.get('UserEmail') , userName=session.get('User') , type=session.get('type'))



@app.route('/Dashboard/furniture/genr')
def fur_gnsea():
    if session.get('UserEmail') == None:
        return redirect('/')
    if session.get('type') not in ['furniture', 'all']:
        return redirect('/')
    fields = [
    {'name': 'office_number', 'label': 'رقم المكتب'},

    ]


    connection = sqlite3.connect('office_inventory.db')
    cursor = connection.cursor()
    cursor.execute('SELECT * FROM office')
    all = cursor.fetchall()



    connection = sqlite3.connect('office_inventory.db')
    cursor = connection.cursor()
    cursor.execute('SELECT name, desk FROM office')
    cc = cursor.fetchall()
    x = 0
    for i in cc:
        x += int(i[1])



    connection = sqlite3.connect('office_inventory.db')
    cursor = connection.cursor()
    cursor.execute('SELECT * FROM office')
    ttcc = cursor.fetchall()
    dd = 0
    for i in ttcc:
        dd += 1
    return render_template('fur_GENS.html',x=x,fields=fields,all=all,dd=dd , userEmail=session.get('UserEmail') , userName=session.get('User') , type=session.get('type'))



if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8080) 